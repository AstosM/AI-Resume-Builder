import os
import json
import uuid
from io import BytesIO
from flask import Flask, render_template, request, send_file, url_for, flash
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

# Optional OpenAI
try:
    import openai
    OPENAI_AVAILABLE = True
except Exception:
    OPENAI_AVAILABLE = False

# PDF/Docx libs optional
try:
    from weasyprint import HTML
    WEASY_AVAILABLE = True
except Exception:
    WEASY_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

load_dotenv()

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "devsecret")
UPLOAD_FOLDER = os.path.join("static", "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 6 * 1024 * 1024  # 6 MB

# Map template keys to filenames in templates/resume_templates/
TPL_MAP = {
    'template1': 'pro_minimal',
    'template2': 'modern_creative',
    'template3': 'tech_profile',
    'template_exec': 'executive',
    'template_founder': 'founder',
}

def call_openai_for_resume(api_key, structured_input):
    """Calls OpenAI ChatCompletion to get structured JSON. Returns dict or None."""
    if not OPENAI_AVAILABLE or not api_key:
        return None
    try:
        openai.api_key = api_key
        prompt = f"""
You are a helpful assistant that produces JSON for resumes.
Input (JSON):
{json.dumps(structured_input, indent=2)}

Produce valid JSON with keys:
- summary : string (2-3 sentences)
- experiences : list of objects: {{company, role, dates, bullets: [str...] }}
- skills_line: string (comma separated)
Return only JSON.
"""
        resp = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.2,
            max_tokens=700
        )
        text = resp["choices"][0]["message"]["content"]
        parsed = json.loads(text)
        return parsed
    except Exception as e:
        print("OpenAI error:", e)
        return None

def fallback_structure(data):
    """Local fallback that generates a simple structure without AI."""
    summary = f"{data.get('title','')} with experience in {', '.join([s.strip() for s in data.get('skills','').split(',')][:3])}."
    ex_list = []
    raw = data.get("experience_raw","")
    for block in raw.split("||"):
        block = block.strip()
        if not block:
            continue
        parts = [p.strip() for p in block.split("—")]
        company = parts[0] if len(parts)>0 else ""
        role = parts[1] if len(parts)>1 else ""
        dates = parts[2] if len(parts)>2 else ""
        desc = parts[3] if len(parts)>3 else ""
        bullets = []
        if desc:
            bullets.append(desc if len(desc)<=120 else desc[:117]+"...")
        else:
            bullets.append(f"Served as {role} at {company}")
        bullets += ["Delivered key results", "Collaborated across teams"]
        ex_list.append({"company": company, "role": role, "dates": dates, "bullets": bullets})
    skills_line = ', '.join([s.strip() for s in data.get('skills','').split(',') if s.strip()])
    return {"summary": summary, "experiences": ex_list, "skills_line": skills_line}

def render_resume_html(chosen_template, form_data, ai_struct, profile_image_url=None):
    tpl = TPL_MAP.get(chosen_template, chosen_template)
    return render_template(f"resume_templates/{tpl}.html", form=form_data, ai=ai_struct, profile_image=profile_image_url)

def generate_docx_bytes(form_data, ai_struct, profile_image_path=None):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")
    doc = Document()
    doc.add_heading(form_data.get("name",""), level=0)
    doc.add_paragraph(form_data.get("title",""))
    if profile_image_path and os.path.exists(profile_image_path):
        try:
            doc.add_picture(profile_image_path, width=Inches(1.2))
        except Exception:
            pass
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(ai_struct.get("summary",""))
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(ai_struct.get("skills_line",""))
    doc.add_heading("Experience", level=1)
    for ex in ai_struct.get("experiences",[]):
        doc.add_heading(f"{ex.get('role','')} — {ex.get('company','')}", level=2)
        for b in ex.get("bullets",[]):
            doc.add_paragraph(b, style='List Bullet')
    doc.add_heading("Education", level=1)
    doc.add_paragraph(form_data.get("education",""))
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def reportlab_pdf_bytes(lines, profile_image_path=None):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    x, y = 50, 750
    p.setFont("Helvetica-Bold", 18)
    p.drawString(x, y, lines.get("title_line","Resume"))
    y -= 24
    p.setFont("Helvetica", 11)
    for line in lines.get("lines",[]):
        p.drawString(x, y, line[:95])
        y -= 14
        if y < 80:
            p.showPage()
            y = 750
    p.save()
    buffer.seek(0)
    return buffer

@app.route("/", methods=["GET","POST"])
def index():
    if request.method == "POST":
        # collect fields
        form = {
            "name": request.form.get("name","").strip(),
            "title": request.form.get("title","").strip(),
            "email": request.form.get("email","").strip(),
            "phone": request.form.get("phone","").strip(),
            "linkedin": request.form.get("linkedin","").strip(),
            "location": request.form.get("location","").strip(),
            "skills": request.form.get("skills","").strip(),
            "achievements": request.form.get("achievements","").strip(),
            "experience_raw": request.form.get("experience","").strip(),
            "education": request.form.get("education","").strip(),
        }
        apikey = request.form.get("apikey","").strip()
        template_choice = request.form.get("template_choice","template1")

        # handle upload
        uploaded = request.files.get("profile_image")
        profile_image_url = None
        profile_image_path = None
        if uploaded and uploaded.filename:
            filename = secure_filename(uploaded.filename)
            unique = f"{uuid.uuid4().hex}_{filename}"
            saved_path = os.path.join(app.config["UPLOAD_FOLDER"], unique)
            uploaded.save(saved_path)
            profile_image_path = saved_path
            profile_image_url = url_for('static', filename=f"uploads/{unique}")

        # call OpenAI or fallback
        ai_struct = None
        if apikey and OPENAI_AVAILABLE:
            ai_struct = call_openai_for_resume(apikey, form)
        if not ai_struct:
            ai_struct = fallback_structure(form)

        resume_html = render_resume_html(template_choice, form, ai_struct, profile_image_url)
        return render_template("result.html",
                               resume_html=resume_html,
                               form_data=form,
                               ai_struct=ai_struct,
                               template_choice=template_choice,
                               profile_image_path=profile_image_path,
                               profile_image_url=profile_image_url)
    return render_template("index.html")

@app.route("/download/pdf", methods=["POST"])
def download_pdf():
    form_data = json.loads(request.form.get("form_data_json"))
    ai_struct = json.loads(request.form.get("ai_json"))
    template_choice = request.form.get("template_choice")
    profile_image_path = request.form.get("profile_image_path") or None

    html = render_resume_html(template_choice, form_data, ai_struct, profile_image_path)

    # WeasyPrint preferred
    if WEASY_AVAILABLE:
        try:
            pdf = HTML(string=html).write_pdf()
            return send_file(BytesIO(pdf), as_attachment=True, download_name=f"{form_data.get('name','resume')}.pdf", mimetype="application/pdf")
        except Exception as e:
            print("WeasyPrint failed:", e)

    # fallback: simple ReportLab
    lines = {"title_line": form_data.get("name",""), "lines": [ai_struct.get("summary",""), ai_struct.get("skills_line","")]}
    buf = reportlab_pdf_bytes(lines, profile_image_path)
    return send_file(buf, as_attachment=True, download_name=f"{form_data.get('name','resume')}.pdf", mimetype="application/pdf")

@app.route("/download/docx", methods=["POST"])
def download_docx():
    form_data = json.loads(request.form.get("form_data_json"))
    ai_struct = json.loads(request.form.get("ai_json"))
    profile_image_path = request.form.get("profile_image_path") or None

    if DOCX_AVAILABLE:
        docx_buf = generate_docx_bytes(form_data, ai_struct, profile_image_path)
        return send_file(docx_buf, as_attachment=True, download_name=f"{form_data.get('name','resume')}.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    # fallback text
    txt = f"""\
{form_data.get('name', '')}
{form_data.get('title', '')}

{ai_struct.get('summary', '')}

Skills: {ai_struct.get('skills_line', '')}
"""

    return send_file(BytesIO(txt.encode()), as_attachment=True, download_name=f"{form_data.get('name','resume')}.txt", mimetype="text/plain")

if __name__ == "__main__":
    app.run(debug=True)
